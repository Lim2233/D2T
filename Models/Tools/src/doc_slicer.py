"""
docx_semantic_slicer.py

对指定的 .docx 文档进行递归语义切分，每段不超过 1000 字符，
在原路径下生成 <原文件名>_Sliced.json 文件。
"""

import os
import re
import json
import sys
from typing import List, Dict, Any

try:
    from docx import Document
except ImportError:
    print("请先安装 python-docx：pip install python-docx")
    sys.exit(1)


# ---------- 文本提取 ----------
def extract_all_text(docx_path: str) -> List[str]:
    """
    提取 .docx 中所有文本内容，按段落组织（包括表格中的段落）。
    返回段落文本列表，每个元素为一个段落的文本（已去除首尾空白）。
    """
    doc = Document(docx_path)
    paragraphs = []

    def add_para_text(text: str):
        if text and text.strip():
            paragraphs.append(text.strip())

    # 正文段落
    for para in doc.paragraphs:
        add_para_text(para.text)

    # 表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    add_para_text(para.text)

    # 页眉页脚（可选，按需启用）
    for section in doc.sections:
        for para in section.header.paragraphs:
            add_para_text(para.text)
        for para in section.footer.paragraphs:
            add_para_text(para.text)

    return paragraphs


# ---------- 语义切分 ----------
def split_by_sentence(text: str) -> List[str]:
    """
    将文本按中文句号、问号、感叹号进行句子分割，保留分隔符在句尾。
    """
    # 匹配句子结束标点（包括中文和英文标点）
    pattern = r'(?<=[。！？!?])'
    parts = re.split(pattern, text)
    # 去除空字符串
    return [p for p in parts if p.strip()]


def semantic_chunk(text: str, max_len: int = 1000) -> List[str]:
    """
    递归语义切分：若文本长度 ≤ max_len 则直接返回；
    否则按句子切分，并尽量合并句子直到接近 max_len。
    如果单个句子长度超过 max_len，则强制按逗号分割或按长度截断。
    """
    if len(text) <= max_len:
        return [text]

    # 先尝试按句子切分
    sentences = split_by_sentence(text)
    if not sentences:
        # 无句子分隔符，直接按 max_len 截断（可进一步按逗号优化）
        return [text[i:i+max_len] for i in range(0, len(text), max_len)]

    chunks = []
    current_chunk = ""
    for sent in sentences:
        # 如果单个句子本身就超长，进一步按逗号分割或强制截断
        if len(sent) > max_len:
            # 先尝试按逗号分割
            comma_parts = re.split(r'(?<=[，,])', sent)
            if len(comma_parts) == 1:
                # 没有逗号，强制按长度切分
                for i in range(0, len(sent), max_len):
                    chunks.append(sent[i:i+max_len])
            else:
                # 将逗号分割后的子句递归处理（它们仍可能超长）
                for part in comma_parts:
                    if part.strip():
                        chunks.extend(semantic_chunk(part, max_len))
            continue

        # 尝试将句子加入当前块
        if len(current_chunk) + len(sent) <= max_len:
            current_chunk += sent
        else:
            # 当前块已满，保存并开始新块
            if current_chunk:
                chunks.append(current_chunk)
            current_chunk = sent

    if current_chunk:
        chunks.append(current_chunk)

    # 后处理：如果合并后的某些块仍超长（理论上不会，但以防万一）
    final_chunks = []
    for chunk in chunks:
        if len(chunk) > max_len:
            final_chunks.extend(semantic_chunk(chunk, max_len))
        else:
            final_chunks.append(chunk)

    return final_chunks


# ---------- 主流程 ----------
def process_docx(input_path: str, max_len: int = 1000) -> str:
    """
    处理单个 docx 文件，返回生成的 JSON 文件路径。
    """
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"文件不存在: {input_path}")

    # 提取段落文本
    paragraphs = extract_all_text(input_path)

    # 对每个段落进行语义切分
    all_slices = []
    for para_text in paragraphs:
        chunks = semantic_chunk(para_text, max_len)
        all_slices.extend(chunks)

    # 构造输出数据结构
    output_data: Dict[str, Any] = {
        "source": os.path.abspath(input_path),
        "max_slice_length": max_len,
        "slice_count": len(all_slices),
        "slices": [{"id": idx + 1, "text": chunk} for idx, chunk in enumerate(all_slices)]
    }

    # 生成输出路径
    base, _ = os.path.splitext(input_path)
    output_path = f"{base}_Sliced.json"

    # 写入 JSON
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, ensure_ascii=False, indent=2)

    return output_path


# ---------- 命令行入口 ----------
if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("用法: python docx_semantic_slicer.py <待处理的docx文件路径>")
        sys.exit(1)

    docx_file = sys.argv[1]
    try:
        output_file = process_docx(docx_file)
        print(f"切分完成！结果已保存至：{output_file}")
    except Exception as e:
        print(f"处理失败：{e}")
        sys.exit(1)